import streamlit as st
import yt_dlp
import re
import os
from io import BytesIO
from openpyxl import Workbook
from docx import Document

# --------------------------
# 字幕抽出関数
# --------------------------
def extract_subtitles(subtitle_file, keywords, use_censored):
    matches = []
    with open(subtitle_file, "r", encoding="utf-8") as f:
        blocks = f.read().split("\n\n")
    for block in blocks:
        lines = block.strip().split("\n")
        if len(lines) >= 3:
            timestamp = lines[1]
            text = " ".join(lines[2:])
            if keywords:
                for kw in keywords:
                    if kw in text:
                        matches.append((timestamp, text))
                        break
            elif use_censored:
                if "[__]" in text.replace(" ",""):
                    matches.append((timestamp, text))
    return matches

# --------------------------
# 安全なファイル名
# --------------------------
def safe_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)

# --------------------------
# ファイル生成
# --------------------------
def generate_file(matches, title, fmt):
    safe_title = safe_filename(title)
    
    if fmt == "TXT":
        output = "\n".join([f"[{t}] {txt}" for t, txt in matches])
        return BytesIO(output.encode("utf-8")), f"{safe_title}.txt"
    
    elif fmt == "Word":
        doc = Document()
        doc.add_heading(title, level=1)
        for t, txt in matches:
            doc.add_paragraph(f"[{t}] {txt}")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer, f"{safe_title}.docx"
    
    elif fmt == "Excel":
        wb = Workbook()
        ws = wb.active
        ws.append(["Timestamp", "Text"])
        for t, txt in matches:
            ws.append([t, txt])
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer, f"{safe_title}.xlsx"

# --------------------------
# Streamlit UI
# --------------------------
st.title("YouTube 字幕抽出アプリ (ブラウザ版)")

url = st.text_input("YouTube動画URL")
keywords_input = st.text_input("抽出キーワード（カンマ区切り）")
use_censored = st.checkbox("[ __ ]を抽出")
fmt = st.selectbox("保存形式", ["TXT", "Word", "Excel"])

if st.button("抽出開始"):
    if not url or (not keywords_input and not use_censored):
        st.warning("URLとキーワードまたは[ __ ]チェックが必要です")
    else:
        keywords = [kw.strip() for kw in keywords_input.split(",") if kw.strip()] if keywords_input else []
        st.info("📥 字幕抽出中...少々お待ちください")
        
        try:
            # 動画情報取得
            with yt_dlp.YoutubeDL({}) as ydl_info:
                info = ydl_info.extract_info(url, download=False)
            video_title = info.get("title", "untitled")
            st.success(f"🎬 動画タイトル: {video_title}")

            # 字幕ダウンロード
            ydl_opts = {
                "skip_download": True,
                "writesubtitles": True,
                "writeautomaticsub": True,
                "subtitleslangs": ["en"],
                "subtitlesformat": "srt",
                "outtmpl": "subtitle"
            }
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                ydl.download([url])
            
            subtitle_files = [f for f in os.listdir() if f.startswith("subtitle") and f.endswith(".srt")]
            if not subtitle_files:
                st.warning("⚠️ 字幕が見つかりませんでした")
            else:
                subtitle_file = subtitle_files[0]
                matches = extract_subtitles(subtitle_file, keywords, use_censored)
                os.remove(subtitle_file)
                
                if matches:
                    buffer, filename = generate_file(matches, video_title, fmt)
                    st.download_button(
                        label=f"✅ {fmt}ファイルをダウンロード ({len(matches)} 件)",
                        data=buffer,
                        file_name=filename,
                        mime="application/octet-stream"
                    )
                else:
                    st.info("抽出結果は0件でした")
        except Exception as e:
            st.error(f"❌ エラー: {e}")
