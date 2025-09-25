import tkinter as tk
from tkinter import filedialog, messagebox
import yt_dlp
import re
import os
from openpyxl import Workbook
from docx import Document
import threading
import time


# 字幕抽出関数
def extract_subtitles(subtitle_file, keywords, log_callback=None):
    matches = []
    with open(subtitle_file, "r", encoding="utf-8") as f:
        blocks = f.read().split("\n\n")
    for idx, block in enumerate(blocks, 1):
        lines = block.strip().split("\n")
        if len(lines) >= 3:
            timestamp = lines[1]
            text = " ".join(lines[2:])
            # キーワード抽出
            if keywords:
                for kw in keywords:
                    if kw in text:
                        matches.append((timestamp, text))
                        break
            else:
                # キーワード未指定時はチェックONで[ __ ]抽出
                if "[__]" in text.replace(" ",""):
                    matches.append((timestamp, text))
        # ログ更新（任意）
        if log_callback and idx % 5 == 0:  # 5ブロックごとに更新
            log_callback(f"📄 {idx}/{len(blocks)} ブロック処理中...\n")
    return matches


# 安全なファイル名作成

def safe_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)


# 保存処理

def save_matches(matches, title, folder, fmt):
    safe_title = safe_filename(title)
    filepath = os.path.join(folder, f"{safe_title}.{fmt.lower()}")

    if fmt == "TXT":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"=== {title} ===\n")
            for t, txt in matches:
                f.write(f"[{t}] {txt}\n")
            f.write(f"\n合計 {len(matches)} 件\n")
    elif fmt == "Word":
        doc = Document()
        doc.add_heading(title, level=1)
        for t, txt in matches:
            doc.add_paragraph(f"[{t}] {txt}")
        doc.add_paragraph(f"\n合計 {len(matches)} 件")
        doc.save(filepath)
    elif fmt == "Excel":
        wb = Workbook()
        ws = wb.active
        ws.append(["Timestamp", "Text"])
        for t, txt in matches:
            ws.append([t, txt])
        wb.save(filepath)

    return filepath


# GUIログ更新関数

def log_insert(text):
    log_text.insert(tk.END, text)
    log_text.see(tk.END)


# 字幕抽出処理

def run_extraction():
    url = url_entry.get().strip()
    folder = folder_path.get().strip()
    fmt = format_var.get()
    keywords_input = keywords_entry.get().strip()
    use_censored = censored_var.get()

    if not url or not folder or (not keywords_input and not use_censored):
        messagebox.showwarning("入力不足", "URL・保存先・キーワードまたはチェックボックスは必須です")
        return

    keywords = [kw.strip() for kw in keywords_input.split(",") if kw.strip()] if keywords_input else []

    log_text.delete(1.0, tk.END)
    log_text.insert(tk.END, "📥 字幕抽出開始...\n")
    log_text.see(tk.END)

    extracting = True

    def animate():
        dots = ""
        while extracting:
            dots += "."
            if len(dots) > 5:
                dots = ""
            log_text.delete("anim", tk.END)
            log_text.insert(tk.END, f"⏳ 抽出中{dots}\n", "anim")
            log_text.see(tk.END)
            time.sleep(0.5)

    anim_thread = threading.Thread(target=animate)
    anim_thread.start()

    try:
        # 動画情報取得
        with yt_dlp.YoutubeDL({}) as ydl_info:
            info = ydl_info.extract_info(url, download=False)
        video_title = info.get("title", "untitled")
        log_insert(f"🎬 動画タイトル: {video_title}\n")

        # 字幕ダウンロード
        ydl_opts = {
            "skip_download": True,
            "writesubtitles": True,
            "writeautomaticsub": True,
            "subtitleslangs": ["en"],
            "subtitlesformat": "srt",
            "outtmpl": "subtitle"
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])

        subtitle_files = [f for f in os.listdir() if f.startswith("subtitle") and f.endswith(".srt")]
        if not subtitle_files:
            log_insert("⚠️ 字幕が見つかりませんでした\n")
            return
        subtitle_file = subtitle_files[0]
        matches = extract_subtitles(subtitle_file, keywords if not use_censored else [], log_callback=log_insert)
        os.remove(subtitle_file)

        saved_path = save_matches(matches, video_title, folder, fmt)
        log_insert(f"✅ 保存完了: {saved_path} ({len(matches)} 件)\n")
    except Exception as e:
        log_insert(f"❌ エラー: {e}\n")
    finally:
        extracting = False  


# ボタンクリックでスレッド実行

def run_extraction_thread():
    threading.Thread(target=run_extraction).start()


# GUI作成

root = tk.Tk()
root.title("字幕抽出アプリ")

tk.Label(root, text="YouTube動画URL:").grid(row=0, column=0, sticky="e")
url_entry = tk.Entry(root, width=50)
url_entry.grid(row=0, column=1, padx=5, pady=5)

tk.Label(root, text="保存先フォルダ:").grid(row=1, column=0, sticky="e")
folder_path = tk.StringVar()
folder_entry = tk.Entry(root, textvariable=folder_path, width=40)
folder_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
tk.Button(root, text="参照", command=lambda: folder_path.set(filedialog.askdirectory())).grid(row=1, column=2, padx=5, pady=5)

tk.Label(root, text="保存形式:").grid(row=2, column=0, sticky="e")
format_var = tk.StringVar(value="TXT")
tk.OptionMenu(root, format_var, "TXT", "Word", "Excel").grid(row=2, column=1, sticky="w", padx=5, pady=5)

tk.Label(root, text="抽出キーワード（カンマ区切り）:").grid(row=3, column=0, sticky="e")
keywords_entry = tk.Entry(root, width=50)
keywords_entry.grid(row=3, column=1, padx=5, pady=5)

censored_var = tk.BooleanVar()
tk.Checkbutton(root, text="[ __ ]を抽出", variable=censored_var).grid(row=4, column=1, sticky="w", padx=5, pady=5)

tk.Button(root, text="抽出開始", command=run_extraction_thread).grid(row=5, column=1, pady=10)

log_text = tk.Text(root, width=80, height=20)
log_text.grid(row=6, column=0, columnspan=3, padx=10, pady=10)

root.mainloop()
